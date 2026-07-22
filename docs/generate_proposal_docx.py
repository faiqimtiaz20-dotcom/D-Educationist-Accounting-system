"""Generate Word proposal from the markdown content."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).resolve().parent / "DEducationist-Accounting-System-Proposal.docx"


def set_cell_shading(cell, hex_color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tePr = cell._tc.get_or_add_tcPr()
    cell._tc.get_or_add_tcPr().append(shading)


def style_doc(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        hs.font.name = "Calibri"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_rich_para(doc, parts, space_after=6):
    """parts: list of (text, bold)"""
    p = doc.add_paragraph()
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = "Calibri"
        set_cell_shading(hdr[i], "1A365D")

    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
            if r_idx % 2 == 1:
                set_cell_shading(cells[c_idx], "F7FAFC")

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def build():
    doc = Document()
    style_doc(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PROJECT PROPOSAL")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    r.font.name = "Calibri"

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = st.add_run("D’ Educationist Accounting System")
    r2.bold = True
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
    r2.font.name = "Calibri"

    meta = [
        ("Prepared for: ", "D’ Educationist Private Limited"),
        ("Website: ", "https://deducationist.com/"),
        ("Prepared by: ", "______________________________"),
        ("Proposal date: ", "20 July 2026"),
        ("Proposal validity: ", "15 days from the date of issue"),
        ("Reference: ", "DE-ACC-2026-001"),
    ]
    for label, value in meta:
        add_rich_para(doc, [(label, True), (value, False)], space_after=2)

    doc.add_paragraph()

    # 1
    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "This proposal covers the design, development, and delivery of a multi-branch commission accounting system tailored for D’ Educationist’s study-abroad consultancy operations in Pakistan.",
    )
    add_para(
        doc,
        "The system will centralize student records, university commission invoicing, remittances, sub-agent payouts, expenses, payroll, general ledger, approvals, audit trail, and financial reporting — with role-based access for Super Admin, Branch Managers, Accountants, and Counsellors.",
    )
    add_table(
        doc,
        ["Item", "Detail"],
        [
            ["Project fee", "USD 3,500 (fixed)"],
            ["Payment terms", "50% advance · 50% on final delivery"],
            ["Duration", "2 months (8 weeks)"],
            ["Currency", "United States Dollars (USD)"],
            ["Backend", "Node.js"],
            ["Database", "PostgreSQL"],
        ],
        [2.2, 4.5],
    )

    # 2
    add_heading(doc, "2. Client & Project Overview", 1)
    add_table(
        doc,
        ["Field", "Detail"],
        [
            ["Client", "D’ Educationist Private Limited"],
            ["Industry", "Study abroad consultancy & education services"],
            ["Project name", "D’ Educationist Accounting System"],
            [
                "Objective",
                "Digitize and consolidate branch-wise financial and commission operations into one secure web application",
            ],
        ],
        [2.0, 4.7],
    )
    add_heading(doc, "Business challenges addressed", 2)
    add_bullets(
        doc,
        [
            "Scattered Excel / manual tracking of commissions and remittances",
            "Limited branch-wise visibility for management",
            "Manual sub-agent commission and payment reconciliation",
            "Need for counsellor-level and university-level profit visibility",
            "Lack of centralized approvals, audit trail, and role-based access",
        ],
    )

    # 3
    add_heading(doc, "3. Scope of Work", 1)
    add_heading(doc, "3.1 Core modules included", 2)
    add_table(
        doc,
        ["#", "Module", "Description"],
        [
            ["1", "Authentication & security", "Secure login, session handling, role-based access control (RBAC)"],
            ["2", "Dashboard", "Management KPIs and separate counsellor dashboard"],
            ["3", "Master Sheet", "Student / application pipeline with branch and counsellor assignment"],
            ["4", "Invoices", "Multi-currency commission invoices, send / resend, status workflow"],
            ["5", "Remittance", "Receipts against invoices with university, bank, and currency tracking"],
            ["6", "Sub-Agents", "Sub-agent master, commission sheet, payments, and ledger"],
            ["7", "Cash & Expenses", "Petty cash, expenses, bank & cash views"],
            ["8", "Accounting", "General ledger (COA), journal entries, auto-posting foundations"],
            ["9", "Tax & Ledgers", "Tax compliance screens; student, vendor, and sub-agent ledgers"],
            ["10", "Payroll", "Employee payroll with Pakistan salary tax logic"],
            ["11", "Operations", "Documents, approvals workflow, audit trail"],
            ["12", "Reports hub", "Branch, consolidated, commission, tax, and operations reports"],
            ["13", "Settings", "Branches, users & roles, system settings"],
            ["14", "Responsive UI", "Desktop and mobile-friendly layout"],
        ],
        [0.5, 2.0, 4.2],
    )

    add_heading(doc, "3.2 Key reports included", 2)
    add_para(doc, "Branch", bold=True, space_after=2)
    add_bullets(
        doc,
        [
            "Branch Income",
            "Branch Expenses",
            "Branch Profit (with detail drill-down)",
            "Branch Cash Position (with detail drill-down)",
        ],
    )
    add_para(doc, "Consolidated", bold=True, space_after=2)
    add_bullets(
        doc,
        [
            "University-wise P&L",
            "Consolidated P&L",
            "Consolidated Balance Sheet",
            "Consolidated Cash Flow",
        ],
    )
    add_para(doc, "Commission / Operations", bold=True, space_after=2)
    add_bullets(
        doc,
        [
            "Counsellor Report Profit and Loss",
            "Country-wise Report",
            "University-wise Report",
            "Commission earned vs received, sub-agent payout, net margin (as per agreed report set)",
        ],
    )
    add_para(doc, "Tax / Standard", bold=True, space_after=2)
    add_bullets(
        doc,
        ["WHT, GST/SRB, Salary Tax summaries and standard registers (as per agreed list)"],
    )

    add_heading(doc, "3.3 Key business logic included", 2)
    add_bullets(
        doc,
        [
            "Multi-branch data scoping (Super Admin: all branches; others: own branch)",
            "Commission income and remittance outstanding tracking",
            "Sub-agent payable calculation and payment recording",
            "Net profit formula where applicable: Profit/(Loss) − Outstanding = Net Profit",
            "Pakistan payroll tax calculation (as configured)",
            "Audit logging of critical actions",
        ],
    )

    # 4
    add_heading(doc, "4. Technology Stack", 1)
    add_table(
        doc,
        ["Layer", "Technology"],
        [
            ["Frontend", "React, TypeScript, Vite, Tailwind CSS"],
            ["UI components", "Modern component library (Radix / shadcn-style)"],
            ["State", "Zustand (client state) / API-driven data from backend"],
            ["Charts", "Recharts"],
            ["Routing", "React Router"],
            ["Backend", "Node.js (Express / NestJS-style REST API)"],
            ["Database", "PostgreSQL"],
            ["Auth", "JWT / session-based authentication with role-based access"],
            ["Branding", "D’ Educationist name and logo"],
        ],
        [2.0, 4.7],
    )
    add_para(
        doc,
        "Hosting (e.g. Render, Railway, VPS, or client-preferred cloud) will be confirmed at kickoff. Domain, SSL, and server costs remain client-side unless agreed as an add-on.",
        italic=True,
    )

    # 5
    add_heading(doc, "5. Deliverables", 1)
    add_numbered(
        doc,
        [
            "Fully working web application as per Section 3 (React frontend + Node.js backend + PostgreSQL database)",
            "Source code handover (frontend + backend + database migrations/schema) on final payment",
            "Admin / Super Admin access credentials",
            "Basic user guide (PDF or shared document)",
            "One remote training / walkthrough session (up to 2 hours)",
            "Deployment support for first production go-live (hosting credentials to be provided by client, or arranged as optional add-on)",
            "Environment configuration guide (.env.example for frontend and backend)",
        ],
    )

    # 6
    add_heading(doc, "6. Timeline — 2 Months", 1)
    add_table(
        doc,
        ["Phase", "Week", "Activities"],
        [
            [
                "Phase 1 — Kickoff & foundation",
                "Week 1–2",
                "Requirements freeze, branding, Node.js API setup, PostgreSQL schema, auth/RBAC, branches/users, base layout",
            ],
            [
                "Phase 2 — Core operations",
                "Week 3–4",
                "Master Sheet, invoices, remittance, sub-agents, expenses, petty cash",
            ],
            [
                "Phase 3 — Accounting & payroll",
                "Week 5–6",
                "GL/journals, ledgers, payroll, approvals, audit trail, tax screens",
            ],
            [
                "Phase 4 — Reports & UAT",
                "Week 7",
                "Full report suite, filters/sorting, counsellor & university P&L, bug fixes",
            ],
            [
                "Phase 5 — Go-live",
                "Week 8",
                "UAT sign-off, training, deployment, final handover",
            ],
        ],
        [2.2, 1.2, 3.3],
    )
    add_heading(doc, "Milestones", 2)
    add_table(
        doc,
        ["Milestone", "Target", "Payment trigger"],
        [
            ["Project start / kickoff", "Week 1", "Advance 50% due"],
            ["Mid review (core modules demo)", "End of Week 4", "Progress review (no payment)"],
            ["UAT & final delivery", "End of Week 8", "Balance 50% due"],
        ],
        [2.5, 1.8, 2.4],
    )

    # 7
    add_heading(doc, "7. Commercial Terms", 1)
    add_heading(doc, "7.1 Project fee", 2)
    add_table(
        doc,
        ["Description", "Amount (USD)"],
        [
            ["D’ Educationist Accounting System — full development package", "3,500.00"],
            ["Total project value", "USD 3,500.00"],
        ],
        [4.5, 2.2],
    )
    add_heading(doc, "7.2 Payment schedule", 2)
    add_table(
        doc,
        ["Installment", "When", "Amount"],
        [
            ["Advance (50%)", "On proposal acceptance / before development start", "USD 1,750.00"],
            ["Final (50%)", "On UAT sign-off and delivery of source code + go-live build", "USD 1,750.00"],
            ["Total", "", "USD 3,500.00"],
        ],
        [1.8, 3.2, 1.7],
    )
    add_heading(doc, "7.3 Payment method", 2)
    add_para(
        doc,
        "Bank transfer / Wise / Payoneer / other mutually agreed channel. Bank details will be shared on invoice after acceptance.",
    )
    add_heading(doc, "7.4 Inclusions", 2)
    add_bullets(
        doc,
        [
            "Development as per scope in Section 3",
            "Up to two (2) rounds of minor UI/logic revisions during UAT",
            "Training session and basic documentation",
        ],
    )
    add_heading(doc, "7.5 Exclusions (not included in USD 3,500)", 2)
    add_bullets(
        doc,
        [
            "Third-party costs: domain, SSL, hosting, SMS/email provider fees",
            "FBR / PEPPOL / bank API integrations (can be quoted separately)",
            "Mobile native apps (iOS/Android)",
            "Data migration from legacy Excel beyond an agreed sample template",
            "Ongoing monthly support after free warranty period (see Section 8)",
            "Major scope changes after kickoff sign-off",
        ],
    )

    # 8
    add_heading(doc, "8. Warranty & Support", 1)
    add_table(
        doc,
        ["Item", "Terms"],
        [
            ["Bug-fix warranty", "30 days after final delivery for defects in agreed scope"],
            ["Response", "Reasonable effort within 1–2 business days during warranty"],
            [
                "Post-warranty",
                "Optional AMC / support retainer — to be quoted separately (suggested: USD 100–200 / month)",
            ],
        ],
        [2.0, 4.7],
    )

    # 9
    add_heading(doc, "9. Change Requests", 1)
    add_para(
        doc,
        "Any feature not listed in Section 3 will be treated as a change request, estimated in writing (time + cost) and started only after written client approval.",
    )

    # 10
    add_heading(doc, "10. Client Responsibilities", 1)
    add_para(doc, "To keep the 2-month schedule, the client will provide:")
    add_numbered(
        doc,
        [
            "Timely feedback (within 3 business days of each review)",
            "Final logo, brand colors, and letterhead if needed for invoices",
            "Sample data / business rules confirmation (commission rates, branches, roles)",
            "Hosting / domain access (if client-managed) or approval of hosting plan",
            "Named UAT contact for sign-off",
        ],
    )
    add_para(
        doc,
        "Delays in client feedback may extend the delivery date without penalty to the developer.",
    )

    # 11
    add_heading(doc, "11. Assumptions", 1)
    add_bullets(
        doc,
        [
            "Primary users: internal staff (not public student portal)",
            "English UI",
            "Multi-branch Pakistan operations (Karachi, Lahore, Islamabad, etc.)",
            "Multi-currency commission tracking with PKR reporting",
            "Backend stack fixed as Node.js; database fixed as PostgreSQL",
            "One production environment in scope (plus optional staging if hosting allows)",
            "Scope frozen after Week 1 kickoff (except paid change requests)",
        ],
    )

    # 12
    add_heading(doc, "12. Acceptance Criteria", 1)
    add_para(doc, "The project will be considered complete when:")
    add_numbered(
        doc,
        [
            "Modules in Section 3 are available in the production/staging build",
            "Super Admin can manage branches, users, and view all-branch reports",
            "Branch users see only their branch data (as per RBAC)",
            "Counsellor dashboard/report shows counsellor-scoped data",
            "Client completes UAT checklist and provides written / email sign-off",
            "Source code and credentials are handed over after final payment",
        ],
    )

    # 13
    add_heading(doc, "13. Confidentiality & Intellectual Property", 1)
    add_bullets(
        doc,
        [
            "Both parties will keep business and technical information confidential.",
            "Upon full payment, source code and project IP for this system transfer to the client.",
            "Until final payment, the developer retains ownership of the work product.",
            "Third-party libraries remain under their respective open-source licenses.",
        ],
    )

    # 14
    add_heading(doc, "14. Acceptance of Proposal", 1)
    add_para(
        doc,
        "By signing below, the client accepts this proposal, the scope, timeline, and commercial terms.",
    )
    add_table(
        doc,
        ["", "Client", "Developer"],
        [
            ["Name", "", ""],
            ["Designation", "", ""],
            ["Organization", "D’ Educationist Private Limited", ""],
            ["Signature", "", ""],
            ["Date", "", ""],
        ],
        [1.8, 2.5, 2.4],
    )
    add_rich_para(doc, [("Advance due on acceptance: ", True), ("USD 1,750.00", True)])
    add_rich_para(doc, [("Balance due on delivery: ", True), ("USD 1,750.00", True)])
    add_rich_para(doc, [("Total: ", True), ("USD 3,500.00", True)])
    add_rich_para(doc, [("Duration: ", True), ("2 months", True)])

    # 15
    add_heading(doc, "15. Next Steps", 1)
    add_numbered(
        doc,
        [
            "Client reviews and signs this proposal",
            "Advance invoice issued → payment of USD 1,750",
            "Kickoff meeting (Week 1) — requirements freeze & branding",
            "Development as per 8-week plan",
            "Mid demo (Week 4)",
            "UAT, training, go-live (Week 8)",
            "Final payment USD 1,750 → source code handover",
        ],
    )

    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(12)
    r = closing.add_run(
        "Thank you for the opportunity to partner with D’ Educationist. We look forward to delivering a reliable accounting system that supports your multi-branch study-abroad operations."
    )
    r.italic = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    doc.add_paragraph()
    add_rich_para(doc, [("Contact: ", True), ("______________________________", False)])
    add_rich_para(doc, [("Email: ", True), ("______________________________", False)])
    add_rich_para(doc, [("Phone: ", True), ("______________________________", False)])

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
